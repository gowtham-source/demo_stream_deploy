import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(name, p_id, age, bp, chol, cpt, output_file):
    """
    Creates a Word document report with the provided patient details.
    
    Args:
        name (str): Patient name.
        p_id (str): Patient ID.
        age (int): Patient age.
        bp (int): Blood pressure.
        chol (int): Cholesterol level.
        cpt (int): Chest pain type.
        output_file (str): Path to the output Word document.
    
    Returns:
        str: Path to the generated Word document.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Patient Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Patient Name
    doc.add_paragraph(f'Patient Name: {name}',style='Intense Quote')
    
    # Patient ID
    doc.add_paragraph(f'Patient ID: {p_id}',style='Intense Quote')
    
    # Age and Blood Pressure using a table for alignment
    table = doc.add_table(rows=1, cols=2)
    cell1, cell2 = table.rows[0].cells
    cell1.text = f'**Age:** {age}'
    cell2.text = f'**Blood Pressure:** {bp}'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    
    # Cholesterol Level
    doc.add_paragraph(f'**Cholesterol Level:** {chol}')
    
    # Chest Pain Type
    doc.add_paragraph(f'**Chest Pain Type:** {cpt}')
    
    # Save the document
    doc.save(output_file)
    
    return output_file

st.title('Patient Details Form')

name = st.text_input('Patient name', placeholder='Enter the patient name')
p_id = st.text_input('Patient id', placeholder='Enter the patient id')
age = st.number_input('Patient Age', placeholder='Enter the patient Age')
bp = st.number_input('Patient Blood pressure', placeholder='Enter the patient pressure')
chol = st.number_input('Cholestrol Level', placeholder='Enter the Cholestrol Level')
cpt = st.number_input('Chest Pain Type', placeholder='Chest Pain Type')

if st.button('Generate Report'):
    word_file = 'patient_report.docx'
    
    # Create the Word document report
    create_word_report(name, p_id, age, bp, chol, cpt, word_file)
    
    st.success(f'Report generated successfully: {word_file}')
    st.download_button(label='Download Report', data=open(word_file, 'rb'), file_name=word_file)
